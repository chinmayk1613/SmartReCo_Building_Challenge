from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/SmartReco_Recommendation_System_Deep_Dive.docx")
INK = "16182D"
PURPLE = "5445D6"
BLUE = "3F61E8"
LIGHT_PURPLE = "F0EEFF"
LIGHT_BLUE = "EEF3FF"
LIGHT_GRAY = "F3F5F8"
MID_GRAY = "687086"
WHITE = "FFFFFF"
GREEN = "17805C"
AMBER = "A56508"
RED = "A6324A"
CONTENT_DXA = 9360


def set_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_DXA, (widths, sum(widths))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, PURPLE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(str(text))
        set_font(run, size=8.5, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            if len(table.rows) % 2 == 1:
                shade_cell(cells[idx], "F8F9FC")
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx and len(str(value)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            set_font(run, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_font(lead, bold=True, color=INK)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_font(rest, color=INK)
    else:
        run = paragraph.add_run(text)
        set_font(run, color=INK)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_font(run, color=INK)


def add_steps(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.375)
        paragraph.paragraph_format.first_line_indent = Inches(-0.188)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(item)
        set_font(run, color=INK)


def add_callout(doc, title, text, fill=LIGHT_PURPLE, accent=PURPLE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    title_run = paragraph.add_run(title.upper())
    set_font(title_run, size=8.5, bold=True, color=accent)
    detail = cell.add_paragraph()
    detail.paragraph_format.space_after = Pt(0)
    detail.paragraph_format.line_spacing = 1.15
    set_font(detail.add_run(text), size=10, color=INK)
    set_table_geometry(table, [CONTENT_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F5F6FA")
    set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    set_font(paragraph.add_run(text), name="Consolas", size=9, color=INK)
    set_table_geometry(table, [CONTENT_DXA])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, PURPLE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "29356A", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run("SMARTRECO  /  RECOMMENDATION SYSTEM"), size=8, bold=True, color=MID_GRAY)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(fp.add_run("SmartReco technical guide  |  "), size=8, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run = fp.add_run()
    run._r.extend([fld_char1, instr, fld_char2])


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_font(kicker.add_run("TECHNICAL DEEP DIVE"), size=10, bold=True, color=PURPLE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    set_font(title.add_run("How SmartReco Recommends"), size=30, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    set_font(subtitle.add_run("Overall behavioral recommendations and course-detail next steps"), size=15, color=BLUE)

    add_callout(
        doc,
        "Core design principle",
        "Deterministic behavioral evidence and catalog-grounded retrieval select the courses. The LLM explains the already-selected courses persuasively; it is not allowed to invent or independently choose catalog products.",
        fill=LIGHT_PURPLE,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(26)
    meta.paragraph_format.space_after = Pt(2)
    set_font(meta.add_run("Implementation-aligned guide  |  7 August 2026"), size=10, bold=True, color=MID_GRAY)
    audience = doc.add_paragraph()
    audience.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(audience.add_run("Audience: judges, architects, engineers, product owners, and operators"), size=9.5, italic=True, color=MID_GRAY)
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    add_heading(doc, "1. Executive summary", 1)
    add_body(doc, "SmartReco has two distinct recommendation experiences that share one governed LangGraph workflow but use different triggers and ranking objectives.")
    add_table(doc, ["Experience", "When it runs", "Primary evidence", "What the user sees"], [
        ["Overall - Recommended for you", "After new behavioral trigger points reach 8", "Recency-decayed user profile plus semantic catalog retrieval", "Stored narrative and up to three behavior-first courses on the home page"],
        ["Course detail - You may also be interested in", "Once per course-page visit or refresh", "Current course as the relevance anchor plus accumulated behavior", "A current-course learning story and up to three defensibly related next courses"],
    ], [1900, 2200, 2800, 2460], 8.3)
    add_callout(doc, "Authoritative recommendation", "The removed 'Your strongest interests so far' widget is not part of this design. The home-page 'Recommended for you' panel is the authoritative overall agentic recommendation.")

    add_heading(doc, "2. End-to-end architecture", 1)
    add_steps(doc, [
        "Observe authenticated, user-scoped activity in the browser without blocking navigation.",
        "Batch and persist idempotent activity events in the application database.",
        "Derive typed behavioral signals and update the private user-interest profile.",
        "Choose the scope: wait for the overall score threshold, or start immediately for a course-page visit.",
        "Run the LangGraph workflow to load context, retrieve and rank catalog products, verify them through the MCP boundary, generate grounded wording, validate it, and persist it.",
        "Display only stored, active recommendations and keep every service call observable through the local invocation ledger and LangSmith traces when enabled.",
    ])
    add_code(doc, "Browser -> event batch -> activity_events -> behavioral_signals -> user_interest_profile\n        -> trigger/scope -> LangGraph -> RAG rank -> MCP verify -> Mesh LLM\n        -> validate -> recommendations + recommendation_items -> learner UI")

    add_heading(doc, "3. Activity collection and event persistence", 1)
    add_body(doc, "The browser tracker queues events in local storage, flushes automatically every 1.5 seconds, flushes immediately at ten queued events, and sends at most 100 events per batch. The request uses keepalive, CSRF protection, an authenticated session, and an event UUID. If delivery fails, the batch is returned to the local queue instead of being silently lost.")
    add_bullets(doc, [
        "Every accepted event is keyed to the authenticated user; recommendations and signals never use another learner's evidence.",
        "The backend de-duplicates by event_id and rejects product references that do not exist.",
        "Personalization can be disabled per user; disabled users create no recommendation events.",
        "Course impressions are tracked efficiently with IntersectionObserver. Active dwell counts only foreground time.",
    ])

    add_heading(doc, "4. Event-to-signal model", 1)
    add_body(doc, "Two numeric concepts must not be confused. Signal strength contributes to the learner's topic profile. Trigger points only decide whether enough new behavior exists to justify another overall LangGraph/LLM run.")
    signal_rows = [
        ["page_viewed", "Browse", "0.05", "0.5", "Weak navigation evidence"],
        ["category_selected", "Topic interest", "0.25", "1.0", "Explicit category exploration"],
        ["search_submitted", "Explicit intent", "0.70", "6.0", "Strong stated intent"],
        ["product_impression", "Exposure", "0.03", "0.0", "Visible course; no run trigger"],
        ["product_viewed", "Product interest", "0.35", "3.0", "Course detail opened"],
        ["product_clicked", "Product interest", "0.55", "5.0", "Deliberate course selection"],
        ["active_dwell", "High engagement", "Dynamic", "4.0 initial", "At least 15 active seconds"],
        ["added_to_cart", "Purchase intent", "0.95", "10.0", "Immediate strong intent"],
        ["cart_viewed", "Cart review", "0.30", "1.5", "Reconsidering saved course"],
        ["removed_from_cart", "Cart released", "0.00", "4.0", "Makes course eligible again"],
        ["recommendation_impression", "Recommendation exposure", "0.00", "0.0", "Measurement only"],
        ["recommendation_clicked", "Recommendation response", "0.75", "6.0", "Positive feedback"],
        ["recommendation_dismissed", "Negative feedback", "-0.85", "6.0", "Persistent Not for me"],
        ["purchase_started", "Purchase intent", "0.90", "10.0", "High commercial intent"],
        ["purchase_completed", "Conversion", "1.00", "12.0", "Completed outcome"],
    ]
    add_table(doc, ["Event", "Signal", "Strength", "Trigger", "Interpretation"], signal_rows, [1900, 1850, 900, 900, 3810], 7.8)

    add_heading(doc, "4.1 Dwell-time behavior", 2)
    add_body(doc, "Dwell is active course-page time, not wall-clock time. Hidden-tab time is not counted. The first checkpoint is created at 15 active seconds, followed by approximately 15-second checkpoints. For the same user, session, and course, later checkpoints update one HIGH_ENGAGEMENT signal instead of flooding the profile.")
    add_code(doc, "dwell_strength = min(0.90, 0.35 + log(1 + min(seconds, 300)) / 10)")
    add_body(doc, "The first qualifying dwell contributes 4 trigger points. Later checkpoints for the same signal contribute 1 point. Dwell below 15 seconds advances the event cursor but creates no signal.")

    add_heading(doc, "5. Behavioral profile, lookback, and 72-hour decay", 1)
    add_body(doc, "The recommender rebuilds the profile from up to the latest 500 signals for that user. There is no currently enforced fixed lookback-day filter. Each signal is assigned a nominal 30-day expires_at value, but the profile query does not yet exclude expired signals. Consequently, the effective lookback is the latest 500 signals with recency decay.")
    add_code(doc, "topic_weight += strength * confidence * 2^(-age_hours / 72)")
    add_table(doc, ["Signal age", "Weight retained", "Meaning"], [
        ["0 hours", "100%", "Fresh evidence"],
        ["24 hours", "79%", "Recent behavior remains strong"],
        ["48 hours", "63%", "Still influential"],
        ["72 hours", "50%", "One half-life"],
        ["144 hours", "25%", "Two half-lives"],
        ["216 hours", "12.5%", "Historical context only"],
    ], [1800, 1800, 5760], 8.8)
    add_bullets(doc, [
        "Primary intent: the highest positive recency-decayed topic weight.",
        "Secondary intents: up to the next three positive topics.",
        "Recent searches: the latest ten search strings; up to five are used in the RAG query and LLM prompt.",
        "Positive products: strong positive signals with strength above 0.5.",
        "Negative products: the current state of explicit Not for me actions; later positive behavior can clear the negative state.",
        "Journey stage: exploration, comparison, purchase_intent, or conversion.",
        "Profile confidence: positive weight from the top three topics, capped at 1.0.",
        "Profile hash: prevents duplicate runs when the meaningful profile content has not changed.",
    ])
    add_callout(doc, "Current temporal limitation", "The profile knows recency through decay but does not explicitly label stable, emerging, shifted, or mixed intent. The LLM receives aggregated primary/secondary interests, not a full chronological event timeline. A future 48-hour / 7-day / 30-day multi-window profile would make interest shifts explicit.", fill="FFF7E9", accent=AMBER)

    add_heading(doc, "6. Overall recommendation lifecycle", 1)
    add_heading(doc, "6.1 When it appears and refreshes", 2)
    add_body(doc, "A new user does not receive the authoritative home-page recommendation after a passive impression. New trigger points must accumulate to the configured threshold of 8.0. Once a recommendation is stored, it remains visible while later evidence accumulates. After each successful claim, trigger_score resets to zero; another 8 points are required for the next refresh.")
    add_table(doc, ["Behavior sequence", "Points", "Run?"], [
        ["Search + product view", "6 + 3 = 9", "Yes"],
        ["Product click + product view", "5 + 3 = 8", "Yes"],
        ["Three product views", "3 + 3 + 3 = 9", "Yes"],
        ["Add to cart", "10", "Yes immediately"],
        ["Sixteen page views", "16 x 0.5 = 8", "Technically yes"],
        ["Product impressions only", "0", "Never"],
    ], [4300, 2300, 2760], 8.8)
    add_bullets(doc, [
        "A 20-second cooldown prevents rapid duplicate overall runs.",
        "An active run is not duplicated; new evidence can request a bounded follow-up run.",
        "At most two automatic catch-up runs are executed after a completed run.",
        "If the profile hash is unchanged, the system returns duplicate_profile instead of spending another LLM call.",
        "The home page is server-rendered, so a newly persisted recommendation appears on the next navigation or refresh.",
    ])

    add_heading(doc, "6.2 Overall RAG retrieval and deterministic ranking", 2)
    add_body(doc, "The profile becomes a retrieval query containing the primary learning goal, secondary interests, up to five recent searches, and journey stage. The vector store returns up to 40 semantic candidates. SQL then verifies the current active catalog and scores every eligible product.")
    add_code(doc, "overall_score = 0.42 * semantic_similarity\n              + 0.25 * behavioral_topic_match\n              + 0.18 * search_term_match\n              + 0.08 * catalog_quality\n              + 0.07 * popularity\n              + 0.08 engagement_bonus (when applicable)")
    add_bullets(doc, [
        "Saved, purchased, explicitly dismissed, and other scope-excluded products are removed before selection.",
        "The semantic score comes from Qdrant. Mesh embeddings are currently disabled, so the vector layer uses the deterministic local embedding fallback.",
        "At most two selected courses come from one category before a broader fallback fills any remaining positions.",
        "The graph requests the top three courses for the learner-facing panel.",
    ])

    add_heading(doc, "6.3 Persistence and replacement", 2)
    add_body(doc, "A successful run stores one Recommendation row, ordered RecommendationItem rows, a profile snapshot, model name, scores, explanation, and seven-day expiry timestamp. Before insertion, the prior active overall recommendation is marked superseded. The current homepage loader does not yet enforce expires_at, so an inactive learner's last active recommendation remains visible until replaced.")

    add_heading(doc, "7. Course-detail recommendation lifecycle", 1)
    add_body(doc, "The course-detail experience is not gated by the overall 8-point score. Opening or refreshing a course page creates a unique visit ID and attempts to claim one course-scoped run. This allows a new learner to receive useful next steps even before a rich behavioral history exists, because the current course itself supplies strong context.")
    add_table(doc, ["Property", "Course-detail behavior"], [
        ["Trigger", "Course page opened; no 8-point threshold"],
        ["Anchor", "The exact current course title, category, level, skills, description, and outcomes"],
        ["Personalization", "Accumulated user profile adjusts related-course order and learning-path direction"],
        ["Provider calls", "Exactly one configured free-model attempt per page visit; deterministic fallback if it fails"],
        ["Maximum output", "Up to three courses; fewer are allowed when relevance is insufficient"],
        ["UI state", "Generating spinner -> current timestamp on success, or a recoverable failed state"],
        ["Persistence", "User ID, recommendation ID, source course, selected course IDs, scores, copy, and run evidence"],
    ], [2200, 7160], 8.7)
    add_callout(doc, "Visit evidence timing", "The course-scoped run is queued during the page GET. The product_viewed event and dwell checkpoints from that same visit are recorded immediately afterward but are intentionally marked for the next visit, avoiding a second LLM workflow during the current visit.")

    add_heading(doc, "7.1 Contextual relevance gate", 2)
    add_body(doc, "A candidate cannot enter the course-detail ranking merely because the learner likes it globally. It must first be defensibly related to the course being viewed through at least one gate:")
    add_bullets(doc, [
        "Same category: path relevance = 1.00.",
        "Catalog-defined adjacent category: path relevance = 0.72.",
        "Shared meaningful skills without a configured category edge: path relevance = 0.58.",
        "No same domain, adjacent path, or shared skill: candidate is rejected before ranking.",
    ])
    add_body(doc, "The catalog currently defines adjacent paths across Agentic AI, LLMs, Generative AI, Python, Java, Scala, Web Technologies, Data Engineering, MLOps, and Cloud & DevOps.")

    add_heading(doc, "7.2 Contextual-behavioral score", 2)
    add_code(doc, "context_score = 0.40 * current_course_semantic_similarity\n              + 0.27 * learning_path_relevance\n              + 0.18 * learner_behavior\n              + 0.10 * shared_skill_overlap\n              + 0.05 * level_progression")
    add_body(doc, "Learner behavior is itself a bounded blend of category affinity (65%), recent-search match (25%), and prior positive product engagement (10%). Level progression rewards sensible moves such as beginner to intermediate or intermediate to advanced.")
    add_code(doc, "fit_confidence = min(0.98,\n    0.35 * semantic + 0.35 * path + 0.15 * behavior\n+    + 0.10 * level_progression + 0.05 * evidence_coverage)")
    add_body(doc, "interest_likelihood is a logistic transformation of the hybrid score: 1 / (1 + exp(-8 * (score - 0.42))). It is an interpretable fit indicator, not a calibrated probability of purchase.")
    add_bullets(doc, [
        "The current course, carted courses, purchased courses, and explicitly dismissed courses are excluded.",
        "At most two results may come from one category, preventing 'the other courses in this department' from filling the entire panel.",
        "The third result must contribute a defensible adjacent learning-path or shared-skill angle.",
    ])

    add_heading(doc, "8. The shared LangGraph workflow", 1)
    add_code(doc, "START -> load -> retrieve -> verify_with_mcp -> generate -> validate\n                                                       | valid   -> persist -> END\n                                                       | invalid -> fallback -> persist -> END")
    graph_rows = [
        ["load", "Load behavioral profile, exclusions, run scope, and current-course context when present.", "profile, query, optional context_product_id"],
        ["retrieve", "Use overall or contextual RAG, apply deterministic eligibility and ranking, select up to three.", "selected candidates + retrieval metrics"],
        ["verify_with_mcp", "Read active SQL catalog details through the governed MCP boundary and preserve exact order/IDs.", "verified product IDs"],
        ["generate", "Call Mesh API for grounded persuasive JSON; overall can fail over across three models, contextual uses one free model.", "headline, narrative, item reasons, usage"],
        ["validate", "Ensure generated IDs exactly match selected IDs with no missing, invented, or duplicate course.", "validation errors or valid route"],
        ["fallback", "Create deterministic grounded wording if validation fails.", "safe copy for exact selected IDs"],
        ["persist", "Supersede the prior active recommendation in the same scope and store recommendation/items/scores.", "durable current recommendation"],
    ]
    add_table(doc, ["Node", "Definition", "Output / relationship"], graph_rows, [1500, 4930, 2930], 8.1)
    add_body(doc, "LangGraph uses an InMemorySaver checkpointer keyed by recommendation run ID. Durable operational status, current node, graph-state summaries, leases, errors, tokens, and outputs are also stored in SQL. The SQL evidence survives process restarts; the in-memory LangGraph checkpoint itself does not.")

    add_heading(doc, "9. RAG, MCP, LLM, and LangChain responsibilities", 1)
    add_table(doc, ["Component", "What it does", "What it must not do"], [
        ["RAG / Qdrant", "Retrieve semantically similar catalog IDs; support overall and current-course queries.", "Return invented courses or bypass active SQL state."],
        ["Deterministic ranker", "Apply behavior, search, path, skill, quality, exclusion, and diversification rules.", "Delegate product selection to persuasive generation."],
        ["MCP", "Expose read-only profile/signal/catalog tools and verify final product IDs against active SQL rows.", "Write recommendations, modify users, or replace the graph."],
        ["Mesh LLM", "Write the headline, narrative, and course-specific reasons for already-ranked products.", "Invent IDs, prices, discounts, outcomes, sensitive traits, or choose unrelated products."],
        ["LangGraph", "Make the workflow explicit, conditional, observable, lease-aware, recoverable, and persistable.", "Act as the model or vector database."],
        ["LangSmith", "Trace graph, retriever, tool, and individual provider-attempt spans when configured.", "Replace the durable local telemetry ledger or recommendation database."],
        ["LangChain", "Used indirectly through LangGraph/LangSmith ecosystem primitives and OpenAI-compatible model integration; no separate LangChain agent chooses products.", "Obscure deterministic business rules behind an unconstrained agent."],
    ], [1500, 4480, 3380], 7.9)

    add_heading(doc, "9.1 Where the LLM is important", 2)
    add_body(doc, "The LLM is important for communication, not catalog truth. It turns structured evidence and verified candidates into a coherent, human-facing story. For a course detail page, it must explicitly connect what the current course teaches, what the candidate adds, and how that combination helps the learner understand or build more.")
    add_bullets(doc, [
        "Input: primary and secondary interests, up to five recent searches, journey stage, optional current-course record, and exact verified candidate records with fit scores.",
        "Output: JSON containing one concise headline, two or three warm narrative sentences, and one reason per exact product ID.",
        "Temperature: 0 for repeatability. Contextual maximum output: 700 tokens. Overall maximum output: 4,000 tokens.",
        "Current contextual model: minimax/m2-her through Mesh API, one attempt per page visit.",
        "Current overall chain: minimax/m2-her -> tencent/hy3 -> openai/gpt-4o-mini, stopping after success or a non-recoverable gateway/authentication failure.",
        "If all eligible provider attempts fail, deterministic provider fallback preserves grounded functionality.",
    ])

    add_heading(doc, "10. Worked example A - changing interests over four days", 1)
    add_body(doc, "Scenario: during days 1-2 the learner searches for Machine Learning and Agentic AI. During days 3-4 the learner searches Web Technologies and performs a smaller amount of Java exploration.")
    add_table(doc, ["Time", "Illustrative behavior", "Trigger effect", "Profile effect"], [
        ["Day 1", "Search ML (6) + view ML course (3)", "9 points; first overall run", "ML becomes primary"],
        ["Day 2", "Search Agentic AI (6) + click course (5)", "11 new points; refresh after cooldown", "Agentic AI competes with ML"],
        ["Day 3", "Search Web Technologies (6) + view course (3)", "9 new points; another refresh", "Recent Web evidence has about 79-100% weight"],
        ["Day 4", "View/click Web (3/5) + one Java view (3)", "Threshold can be crossed again", "Web likely becomes primary; Java may be secondary"],
    ], [1200, 3300, 2150, 2710], 8.2)
    add_body(doc, "Because AI evidence from days 1-2 still retains roughly 50-63% of its original recency weight by day 4, the current implementation may keep AI as a secondary interest. The ranker should prioritize Web when the newer weighted evidence overtakes AI, but the system does not yet explicitly label the pattern as an intent shift.")
    add_callout(doc, "How the LLM interprets it today", "The LLM sees the aggregated primary intent, secondary interests, recent searches, and selected courses. It can explain Web as the current direction, but it does not receive a statement such as 'interest shifted from AI to Web' unless that interpretation is added upstream. The deterministic profile and ranker—not the LLM—should decide whether older AI context is relevant.", fill=LIGHT_BLUE, accent=BLUE)
    add_heading(doc, "Recommended temporal upgrade", 2)
    add_body(doc, "For a production-grade transition detector, calculate separate 48-hour, 7-day, and 30-day profiles and pass an explicit intent_mode of stable, emerging, shifted, or mixed. A defensible starting blend is 65% short-term, 25% seven-day, and 10% thirty-day evidence. Older interests should be fused into the story only when catalog relationships make the connection useful.")

    add_heading(doc, "11. Worked example B - Streaming Data Engineering detail page", 1)
    add_body(doc, "A learner opens Streaming Data Engineering. The contextual run starts immediately and centers that exact course. The vector query includes its event-stream domain, level, skills, and description. A candidate such as Java Development Foundations cannot enter merely because Java is globally popular.")
    add_steps(doc, [
        "Eligibility gate keeps Data Engineering courses, configured adjacent paths such as MLOps/Cloud/Python, or courses sharing meaningful stream-processing skills.",
        "Each eligible candidate receives semantic, learning-path, behavior, shared-skill, and level-progression scores.",
        "The ranker selects up to three, with at most two from the same category.",
        "MCP verifies that the selected IDs still exist and are active in SQL.",
        "The LLM writes a story such as: Streaming Data Engineering gives you event-flow foundations; an MLOps course adds deployment and monitoring capabilities for those pipelines.",
        "Validation rejects any output that changes, duplicates, or invents an ID; the verified recommendation is stored with the source-course ID.",
    ])
    add_callout(doc, "Confidence language", "The displayed fit confidence is evidence-based but not a statistically calibrated probability that the learner will purchase or complete the course. It should be presented as fit confidence, not guaranteed likelihood.", fill="FFF7E9", accent=AMBER)

    add_heading(doc, "12. Exclusions, feedback, and recommendation refresh", 1)
    add_table(doc, ["Behavior", "Selection effect", "Future behavior"], [
        ["Added to cart", "Exact course excluded from recommendations", "Related-course behavior still informs the profile"],
        ["Removed from cart", "Course becomes eligible again", "Does not create persistent negative preference"],
        ["Purchase completed", "Purchased course excluded", "Conversion strengthens the topic/journey evidence"],
        ["Not for me", "Exact product enters the negative set", "Later positive action can clear that negative state"],
        ["Recommendation clicked", "Positive product/topic evidence", "Can help cross the next 8-point refresh threshold"],
    ], [1900, 3550, 3910], 8.2)

    add_heading(doc, "13. Reliability, observability, and recovery", 1)
    add_bullets(doc, [
        "Recommendation runs have queued/running/succeeded/failed states, current node, retry count, lease expiry, and error evidence.",
        "Contextual leases use the contextual model timeout plus 35 seconds; overall runs use a 300-second lease.",
        "A queued contextual run older than 60 seconds, or a run with an expired lease, is closed as stale and safely replaced on the next visit.",
        "Every RAG, MCP, LangGraph, and LLM provider attempt creates local service-invocation evidence with user/run correlation, latency, tokens, model, attempt, failure, failover decision, and provider receipt when available.",
        "LangSmith traces the recommendation chain, RAG retriever span, MCP verification tool span, generation chain, and one LLM span per Mesh provider attempt when tracing is enabled.",
        "Local and LangSmith numbers are reconciled by correlation ID; delayed export is reported rather than force-matched.",
    ])

    add_heading(doc, "14. Current limitations and recommended next changes", 1)
    limitation_rows = [
        ["No fixed enforced signal lookback", "Expired signals can still enter the latest-500 profile query", "Filter expires_at and schedule cleanup; enforce a 30-day maximum"],
        ["No explicit intent-shift state", "Older and newer topics may be blended without a declared transition", "Add 48h/7d/30d profiles and shift confidence"],
        ["Heuristic trigger threshold", "8 points is explainable but not statistically optimized", "Replay history and A/B test CTR, dismissal, conversion, cost, and latency"],
        ["Fit likelihood not calibrated", "A 0-1 score may be mistaken for purchase probability", "Calibrate with outcomes or label strictly as fit confidence"],
        ["In-memory LangGraph checkpoint", "Graph checkpoint does not survive a process restart", "Use a durable LangGraph checkpointer while retaining SQL run evidence"],
        ["Homepage expiry not enforced", "Last active overall recommendation can remain past seven days", "Apply expires_at in the loader or schedule a refresh"],
        ["Local deterministic embeddings", "Semantic quality may be lower than production embeddings", "Enable approved Mesh embedding model and compare retrieval metrics"],
        ["Home-page update needs navigation", "Completed background refresh is not immediately injected", "Add a lightweight recommendation lifecycle poll/SSE update"],
    ]
    add_table(doc, ["Current condition", "Why it matters", "Recommended action"], limitation_rows, [2350, 3270, 3740], 8.0)

    add_heading(doc, "15. Presenter-ready explanation", 1)
    add_callout(doc, "Thirty-second summary", "SmartReco does not ask an LLM to guess courses. It captures private user behavior, converts it into recency-decayed evidence, and waits until meaningful activity justifies an overall refresh. RAG retrieves real catalog candidates, deterministic scoring and exclusions choose the courses, MCP verifies them, and LangGraph controls generation, validation, fallback, and persistence. On a course page, the current course becomes the hard relevance anchor while the learner profile personalizes the next step. The LLM's job is to explain the verified path persuasively and safely.")
    add_heading(doc, "Key distinction", 2)
    add_table(doc, ["Question", "Answer"], [
        ["Who chooses the courses?", "RAG plus deterministic ranking and business rules."],
        ["Who verifies catalog truth?", "The read-only MCP/SQL verification boundary."],
        ["What does the LLM do?", "Writes grounded learner-facing narrative and item reasons."],
        ["What does LangGraph do?", "Controls explicit nodes, edges, conditional validation, fallback, status, and observability."],
        ["Why is this a recommender, not a course suggester?", "Selections depend on user-specific multi-event behavior, temporal decay, exclusions, semantic retrieval, ranking scores, feedback, and stored refresh cycles."],
    ], [2500, 6860], 8.8)

    add_heading(doc, "Appendix A - Implementation source map", 1)
    add_table(doc, ["Concern", "Primary implementation"], [
        ["Browser event batching and dwell", "app/static/js/tracker.js"],
        ["Event ingestion and learner routes", "app/routes.py"],
        ["Signal rules, 72-hour decay, profile", "app/services/signals.py"],
        ["Overall/contextual retrieval and LangGraph", "app/services/recommendation.py"],
        ["Mesh prompts, LLM, embedding fallback", "app/services/mesh.py"],
        ["Read-only MCP tools", "app/mcp_server.py and app/services/mcp_catalog.py"],
        ["Vector synchronization and search", "app/services/vector_store.py"],
        ["Durable schema", "app/models.py"],
        ["Local telemetry and LangSmith correlation", "app/services/observability.py and app/services/langsmith_reconciliation.py"],
        ["Runtime settings", "app/config.py and .env"],
    ], [3400, 5960], 8.7)

    add_heading(doc, "Appendix B - Configuration values documented here", 1)
    add_table(doc, ["Setting", "Current value", "Purpose"], [
        ["recommendation_min_trigger_score", "8.0", "Overall refresh threshold"],
        ["recommendation_cooldown_seconds", "20", "Prevent rapid duplicate overall calls"],
        ["signal profile limit", "500 latest", "Upper bound per profile rebuild"],
        ["recency half-life", "72 hours", "Old interests decay gradually"],
        ["nominal signal expiry", "30 days", "Assigned but not yet enforced in profile query"],
        ["recent search limit", "10 stored / 5 prompted", "Current intent evidence"],
        ["overall semantic candidate limit", "40", "RAG candidate pool"],
        ["learner-facing output", "Up to 3 courses", "Concise recommendation panel"],
        ["contextual LLM timeout", "25 seconds", "User-facing page-visit bound"],
        ["overall LLM timeout", "90 seconds", "Background failover allowance"],
        ["Mesh model mode", "free", "Uses minimax/m2-her first"],
        ["Mesh embeddings", "disabled", "Uses local deterministic embeddings"],
    ], [3100, 1800, 4460], 8.4)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
